"""backend/v5/docx_gen.py

ГЕНЕРАЦИЯ WORD (DOCX) В ФОРМАТЕ КАК В КОРПОРАТИВНОМ ШАБЛОНЕ

Требование: ЛТ должен быть 1-в-1 по вёрстке со шаблоном (test.docx).

ВАЖНОЕ ПРО ШАБЛОН
  Шаблон (lt_template.docx) почти полностью сделан таблицами, и внутри
  главной таблицы есть ВЛОЖЕННАЯ таблица (nested table), где находятся:
    - 1–4 блоки (Описание потребности / Продукты / Требование / Реализация)
    - юридическая «портянка» (её не трогаем)
    - 8. Подробный список задач (таблица задач)

Поэтому единственный надёжный способ получить «точь-в-точь»:
  1) Открыть шаблон как основу документа (Document(template)).
  2) Точечно изменить значения в ячейках (не ломая рамки/мерджи/стили).
  3) В таблице задач: удалить старые строки задач и вставить свои,
     клонируя XML строки, чтобы сохранить стиль.

ПРАВИЛА ДЛЯ ТАБЛИЦЫ ЗАДАЧ (по пожеланию пользователя)
  - min = max
  - стоимость распределяется ПРОПОРЦИОНАЛЬНО баллам (pts) по serviceItems
  - лицензии (licItems) идут отдельными строками фиксированной суммой
  - строка «Проверка заказчиком» всегда есть: 4950 / 4950, длительность = как в шаблоне
  - в строке «Итого» длительность НЕ трогаем (оставляем как в шаблоне)

"""

from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.table import Table

from .models import V5Input, V5Result


TEMPLATE_PATH = Path(__file__).resolve().parent / "templates" / "lt_template.docx"

# Фиксированная строка по требованию пользователя
CUSTOMER_CHECK_RUB = 4950


def _force_arial_8(doc: Document) -> None:
    """Приводит весь документ к шрифту Arial 8.

    Требование: «весь шрифт Arial 8». В docx это означает:
      1) Поменять базовый стиль Normal (на будущее).
      2) Проставить Arial 8 во всех run'ах, включая таблицы и колонтитулы.

    Это чуть «в лоб», но зато надёжно и не зависит от того,
    какие стили/шрифты были внутри шаблона.
    """

    def set_run(run) -> None:
        try:
            run.font.name = "Arial"
            run.font.size = Pt(8)
        except Exception:
            pass

    def walk_paragraphs(paragraphs) -> None:
        for p in paragraphs:
            for r in p.runs:
                set_run(r)

    def walk_tables(tables) -> None:
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    walk_paragraphs(cell.paragraphs)
                    walk_tables(cell.tables)

    # 1) Базовый стиль
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(8)
    except Exception:
        pass

    # 2) Основной текст
    walk_paragraphs(doc.paragraphs)
    walk_tables(doc.tables)

    # 3) Колонтитулы (логотипы/подвалы тоже бывают таблицами)
    for s in doc.sections:
        walk_paragraphs(s.header.paragraphs)
        walk_tables(s.header.tables)
        walk_paragraphs(s.footer.paragraphs)
        walk_tables(s.footer.tables)


def _safe_filename(name: str) -> str:
    name = (name or "").strip() or "LT"
    for ch in '<>:\\"/|?*':
        name = name.replace(ch, "_")
    return " ".join(name.split())[:80] or "LT"


def suggest_filename(inp: V5Input) -> str:
    c = inp.contacts
    base = "Протокол_и_ЛТ"
    if c and c.legal_name:
        base = f"ЛТ_{_safe_filename(c.legal_name)}"
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{base}_{stamp}.docx"


# ---------------------------------------------------------------------------
# Поиск нужных таблиц/ячеек в шаблоне
# ---------------------------------------------------------------------------


def _norm(s: str) -> str:
    return " ".join((s or "").split())


def _find_main_table(doc: Document) -> Table:
    # В шаблоне это первая (и единственная) таблица, но на всякий случай ищем по тексту.
    for t in doc.tables:
        big = "\n".join(_norm(c.text) for r in t.rows for c in r.cells)
        if "Лист Требований" in big and "Заказчик" in big and "Стоимость" in big:
            return t
    if doc.tables:
        return doc.tables[0]
    raise RuntimeError("В шаблоне нет таблиц — неверный lt_template.docx")


def _iter_nested_tables(cell) -> list[Table]:
    """Возвращает все вложенные таблицы внутри ячейки."""
    tables: list[Table] = []
    seen = set()
    for tbl_el in cell._tc.xpath('.//w:tbl'):
        # Один и тот же tbl_el может встречаться несколько раз из-за мерджей.
        key = id(tbl_el)
        if key in seen:
            continue
        seen.add(key)
        tables.append(Table(tbl_el, cell))
    return tables


def _find_nested_lt_table(main: Table) -> Table:
    """Ищем вложенную таблицу, где есть и 'Описание потребности', и 'Подробный список задач'."""
    for row in main.rows:
        for cell in row.cells:
            for t in _iter_nested_tables(cell):
                big = "\n".join(_norm(c.text) for r in t.rows for c in r.cells)
                if "Описание потребности" in big and "Подробный список задач" in big and "Стоимость минимальная" in big:
                    return t
    raise RuntimeError("Не нашёл вложенную таблицу ЛТ (nested) — шаблон отличается")


def _set_value_next_to_label(tbl: Table, label_substr: str, value: str) -> bool:
    """Находит ячейку с label_substr и пишет value в следующую ячейку этой строки."""
    label_substr = _norm(label_substr)
    for r in tbl.rows:
        cells = r.cells
        for i, c in enumerate(cells):
            if label_substr in _norm(c.text):
                # Ищем ближайшую "не такую же" ячейку справа.
                for j in range(i + 1, len(cells)):
                    if _norm(cells[j].text) != _norm(c.text):
                        cells[j].text = value
                        return True
                # Если всё смерджено — пишем прямо в эту же ячейку (лучше так, чем никак).
                c.text = value
                return True
    return False


def _cell_set_text_keep_style(cell, text: str) -> None:
    """Очищаем параграфы в ячейке и записываем новый текст (стили ячейки не трогаем)."""
    # Удаляем все параграфы
    cell.text = ""
    # Теперь можно добавить текст
    cell.text = text


# ---------------------------------------------------------------------------
# Протокол: вставка в начало документа
# ---------------------------------------------------------------------------


def _insert_elements_at_start(doc: Document, elements: list) -> None:
    """Вставляет XML-элементы в начало body (в заданном порядке)."""
    body = doc._body._element
    # вставляем по очереди, увеличивая индекс
    for idx, el in enumerate(elements):
        body.insert(idx, el)


def _build_protocol_at_end(doc: Document, inp: V5Input, res: V5Result) -> list:
    """Создаёт протокол в КОНЦЕ документа и возвращает список XML-элементов,
    которые надо перенести в начало."""
    title = doc.add_paragraph("Протокол заполнения")

    t = doc.add_table(rows=0, cols=2)
    # В разных шаблонах стиль может называться по-разному.
    # Если "Table Grid" отсутствует — просто оставляем дефолтный стиль.
    try:
        t.style = "Table Grid"
    except Exception:
        pass

    def add(k: str, v: str):
        row = t.add_row().cells
        row[0].text = k
        row[1].text = v

    c = inp.contacts
    add("Юрлицо", (c.legal_name if c and c.legal_name else "—"))
    if c and c.inn:
        add("ИНН", c.inn)
    add("Контакт", (c.contact_name if c and c.contact_name else "—"))
    if c and c.phone:
        add("Телефон", c.phone)
    if c and c.email:
        add("Email", c.email)
    if c and c.desired_result:
        add("Желаемый результат", c.desired_result)

    add("Сегменты", ", ".join(inp.segments) if inp.segments else "—")

    # 1С
    if getattr(inp, "onec", None) and inp.onec:
        add("1С конфигурация", inp.onec.config or "—")
        add("1С актуальная", "Да" if inp.onec.actual else "Нет")
    add("Пакет", getattr(res.package, "name", "—"))
    add("Юрлиц", str(inp.org_count or 1))
    add("ККТ", str(len(inp.kkt or [])))

    if inp.kkt:
        add("ККТ список", "; ".join([f"{k.vendor} {k.model}" for k in inp.kkt]))

    # Устройства
    devs = inp.devices or []
    sc = sum(1 for d in devs if d.type == "scanner")
    tsd = sum(1 for d in devs if d.type == "tsd")
    add("Сканеры", str(sc))
    add("ТСД", str(tsd))
    add("Клеверенс (коллективная)", "да" if inp.tsd_collective else "нет")

    if inp.product and inp.product.categories:
        add("Категории ЧЗ", ", ".join(inp.product.categories))

    add("Итого", f"{int(res.costs.total_rub or 0)} ₽")

    # Разрыв страницы
    br = doc.add_paragraph()
    br.add_run().add_break(WD_BREAK.PAGE)

    # Возвращаем элементы (в нужном порядке)
    return [title._element, t._element, br._element]


# ---------------------------------------------------------------------------
# Таблица задач: вставка строк, распределение стоимости
# ---------------------------------------------------------------------------


def _row_contains(row, needle: str) -> bool:
    needle = _norm(needle)
    return any(needle in _norm(c.text) for c in row.cells)


def _find_row_index(tbl: Table, needle: str) -> int:
    for i, r in enumerate(tbl.rows):
        if _row_contains(r, needle):
            return i
    raise RuntimeError(f"Не нашёл строку '{needle}' в вложенной таблице")


def _remove_rows(tbl: Table, start: int, end_exclusive: int) -> None:
    """Удаляет строки [start, end_exclusive)."""
    for _ in range(end_exclusive - start):
        tr = tbl.rows[start]._tr
        tbl._tbl.remove(tr)


def _clone_row_xml(row) -> object:
    return deepcopy(row._tr)


def _insert_row_before(tbl: Table, before_idx: int, tr_xml) -> None:
    tbl._tbl.insert(before_idx, tr_xml)


def _fmt_rub(x: int) -> str:
    # В шаблоне без пробелов/разделителей часто проще, но можно с пробелами.
    return str(int(x))


def _distribute_by_pts(total: int, items: list[tuple[str, int]]) -> list[tuple[str, int]]:
    """Распределяет total по items (label, pts), возвращает (label, rub)."""
    pts_sum = sum(max(0, p) for _, p in items) or 0
    if pts_sum <= 0 or total <= 0:
        return [(lbl, 0) for lbl, _ in items]

    raw = [total * p / pts_sum for _, p in items]
    base = [int(x) for x in raw]
    remainder = total - sum(base)

    # Ранжируем по дробной части
    frac_order = sorted(range(len(items)), key=lambda i: (raw[i] - base[i]), reverse=True)
    for i in frac_order[: max(0, remainder)]:
        base[i] += 1

    return [(items[i][0], base[i]) for i in range(len(items))]


def _parse_pkg_inc(inc: str) -> tuple[list[tuple[str, int]], list[str]]:
    """Разбирает поле package.inc в список строк.

    В data.json это обычно строка вида:
      "Регистрация ЧЗ (4,950); Интеграция ЧЗ (9,900); 1 час настройки"

    Возвращаем:
      - fixed: [(label, rub)] — строки с явной суммой в скобках
      - flex:  [label]        — строки без суммы (их потом распределим по баллам)

    Почему так:
      - Пользователь хочет видеть в Word «что входит» и при этом сумма в таблице
        должна сходиться с итогом.
      - Если в inc есть цена — считаем её фиксированной частью.
      - Остальное распределяем вместе с serviceItems по баллам.
    """
    import re

    inc = (inc or "").strip()
    if not inc:
        return [], []

    # Делим по ';' и '+' — в реальных строках встречается и то, и другое.
    parts = []
    for chunk in re.split(r"[;]+", inc):
        chunk = chunk.strip()
        if not chunk:
            continue
        # Внутри иногда есть "+" ("... + ..."), тоже дробим.
        for p in chunk.split("+"):
            p = p.strip()
            if p:
                parts.append(p)

    fixed: list[tuple[str, int]] = []
    flex: list[str] = []

    # Ищем (4,950) / (4950) / (9 900)
    price_re = re.compile(r"\((\s*[0-9][0-9\s,\.]*?)\s*\)")
    for p in parts:
        m = price_re.search(p)
        if m:
            raw = m.group(1)
            rub = int(re.sub(r"[^0-9]", "", raw) or "0")
            label = price_re.sub("", p).strip()  # убираем (price)
            label = label.rstrip("-–— ")
            fixed.append((label or p, rub))
        else:
            flex.append(p)

    return fixed, flex


def _fill_tasks_table(nested: Table, res: V5Result) -> None:
    """Заполняет блок '8. Подробный список задач' во вложенной таблице."""
    header_idx = _find_row_index(nested, "Стоимость минимальная")  # шапка
    check_idx = _find_row_index(nested, "Проверка заказчиком")
    total_idx = _find_row_index(nested, "Итого")

    # Вставлять будем строки задач между header и check
    first_task_idx = header_idx + 1

    # Шаблон строки для копирования: если в шаблоне есть хотя бы одна строка задач — берём её,
    # иначе клонируем строку проверки заказчиком.
    sample_row = nested.rows[first_task_idx] if first_task_idx < check_idx else nested.rows[check_idx]
    sample_tr = _clone_row_xml(sample_row)

    # Удаляем существующие строки задач (между header и check)
    if first_task_idx < check_idx:
        _remove_rows(nested, first_task_idx, check_idx)
        # после удаления индекс check сдвинулся
        check_idx = _find_row_index(nested, "Проверка заказчиком")
        total_idx = _find_row_index(nested, "Итого")

    # --- Собираем строки задач ---
    # Пользователь ожидает, что таблица задач отражает:
    #   1) базовый состав выбранного пакета (package.inc)
    #   2) выбранные доп.опции (calc.serviceItems)
    #   3) лицензии (calc.licItems)
    #   4) + фиксированная строка «Проверка заказчиком»
    # и при этом сумма должна сходиться с итогом.
    total = int(res.costs.total_rub or 0)

    # 1) Лицензии (фикс в рублях)
    lic_items = [(it.label, int(it.rub or 0)) for it in (res.calc.licItems or []) if int(it.rub or 0) > 0]
    lic_total = sum(r for _, r in lic_items)

    # 2) База пакета: часть строк может иметь явную цену в inc (в скобках)
    pkg_inc = getattr(res.package, "inc", "") or ""
    inc_fixed, inc_flex = _parse_pkg_inc(pkg_inc)
    inc_fixed_total = sum(r for _, r in inc_fixed)

    # 3) Доп.услуги (баллы)
    svc_items = [(it.label, int(it.pts or 0)) for it in (res.calc.serviceItems or []) if int(it.pts or 0) > 0]

    # 4) Строки без явной цены из inc — распределяем как «условные баллы»
    # Чтобы не спорить о весах, даём 1 балл на строку.
    inc_flex_pts = [(lbl, 1) for lbl in inc_flex]

    # 5) Фиксированная строка проверки заказчиком
    check_rub = CUSTOMER_CHECK_RUB

    fixed_total = check_rub + lic_total + inc_fixed_total
    remaining_for_distribution = total - fixed_total
    if remaining_for_distribution < 0:
        remaining_for_distribution = 0

    # Распределяем остаток по (inc_flex + svc_items) по баллам
    distributed = _distribute_by_pts(remaining_for_distribution, inc_flex_pts + svc_items)

    # Итоговый список строк: фикс из inc → распределённые → лицензии
    task_rows: list[tuple[str, int]] = []
    task_rows.extend(inc_fixed)
    task_rows.extend(distributed)
    task_rows.extend(lic_items)

    # Подгоняем последний элемент (кроме проверки заказчиком), чтобы сумма совпала
    cur_sum = sum(r for _, r in task_rows) + check_rub
    delta = total - cur_sum
    if task_rows and delta != 0:
        lbl, rub = task_rows[-1]
        task_rows[-1] = (lbl, max(0, rub + delta))

    # --- Вставляем строки задач ---
    insert_at = check_idx  # перед строкой "Проверка заказчиком"
    for label, rub in task_rows:
        tr = deepcopy(sample_tr)
        _insert_row_before(nested, insert_at, tr)
        # после вставки строка теперь доступна как nested.rows[insert_at]
        r = nested.rows[insert_at].cells

        # Колонки в шаблоне (6):
        # 0-1: Задача (иногда 0 и 1 мерджены)
        # 2: Стоимость минимальная
        # 3: Стоимость максимальная
        # 4: Длительность
        # 5: Комментарий
        r[0].text = label
        r[1].text = label  # из-за мерджа Word часто дублирует — оставляем одинаково
        r[2].text = _fmt_rub(rub)
        r[3].text = _fmt_rub(rub)
        r[4].text = ""  # длительность пустая
        r[5].text = ""

        insert_at += 1

    # --- Строка "Проверка заказчиком" ---
    check_idx = _find_row_index(nested, "Проверка заказчиком")
    check_cells = nested.rows[check_idx].cells
    check_cells[2].text = _fmt_rub(check_rub)
    check_cells[3].text = _fmt_rub(check_rub)
    # check_cells[4] (дни) НЕ трогаем (в шаблоне "3")

    # --- Строка "Итого" ---
    total_idx = _find_row_index(nested, "Итого")
    total_cells = nested.rows[total_idx].cells
    # В "Итого" цена обычно в столбцах 2 и 3
    total_cells[2].text = _fmt_rub(total)
    total_cells[3].text = _fmt_rub(total)
    # total_cells[4] (длительность 11-14) НЕ трогаем


# ---------------------------------------------------------------------------
# Основная сборка
# ---------------------------------------------------------------------------


def build_docx_bytes(inp: V5Input, res: V5Result) -> bytes:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Не найден шаблон: {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)

    # 1) Заполняем значения в ЛТ (в таблицах шаблона)
    main = _find_main_table(doc)
    nested = _find_nested_lt_table(main)

    c = inp.contacts
    legal = c.legal_name if c and c.legal_name else "—"
    contact = c.contact_name if c and c.contact_name else "—"

    # Шапка: меняем номер/дату (в ячейке где "Лист Требований №")
    # Ищем строку/ячейку по подстроке.
    lt_num = "1"  # как в образце
    lt_date = datetime.now().strftime("%d.%m.%Y")
    # В главной таблице обычно это первая строка/первая ячейка.
    main.rows[0].cells[0].text = f"Лист Требований № {lt_num} от {lt_date}"

    _set_value_next_to_label(main, "Заказчик:", legal)
    _set_value_next_to_label(main, "Контактное лицо", contact)

    total = int(res.costs.total_rub or 0)
    _set_value_next_to_label(main, "Стоимость, рубли", f"{total}-{total}")

    # 2) Блоки 1–4 во вложенной таблице
    # Формируем текст из состояния (минимально полезный и короткий)
    seg = ", ".join(inp.segments) if inp.segments else "—"
    pkg = getattr(res.package, "name", "—")
    kkt_list = "; ".join([f"{k.vendor} {k.model}" for k in (inp.kkt or [])]) or "—"

    devs = inp.devices or []
    sc = sum(1 for d in devs if d.type == "scanner")
    tsd = sum(1 for d in devs if d.type == "tsd")

    categories = ", ".join(inp.product.categories) if (inp.product and inp.product.categories) else "—"
    desired = c.desired_result if c and c.desired_result else "—"

    desc_need = (
        f"Сегменты: {seg}. Пакет: {pkg}.\n"
        f"Юрлиц: {inp.org_count or 1}.\n"
        f"Желаемый результат: {desired}"
    )

    if getattr(inp, "onec", None) and inp.onec and inp.onec.config:
        onec = f"{inp.onec.config} ({'актуальная' if inp.onec.actual else 'неактуальная'})"
    else:
        onec = "—"

    products = (
        "Используемые компоненты: 1С/контур учёта, Честный ЗНАК, ККТ/кассовое ПО, сканеры/ТСД, ЭДО (при необходимости).\n"
        f"1С: {onec}. "
        f"ККТ: {kkt_list}. Сканеры: {sc}. ТСД: {tsd}. Категории ЧЗ: {categories}."
    )

    requirement = "Выполнить настройку и проверку сценариев маркировки согласно выбранным пунктам чек‑листа."
    if inp.tsd_collective and tsd:
        requirement += " Включить коллективную работу Клеверенс на ТСД."
    if inp.has_edo:
        requirement += " Обеспечить работу ЭДО."
    if inp.needs_rr:
        requirement += " Настроить разрешительный режим."
    if inp.needs_aggregation:
        requirement += " Настроить агрегацию/КИТУ."

    realization = "Работы выполняются по этапам: диагностика → настройка → тесты → сдача/проверка заказчиком."

    # В шаблоне строки 0–3 соответствуют 1–4 разделам.
    # В каждой строке первый столбец пустой/нумерация, дальше — заголовок, далее поле ввода.
    def write_block(row_idx: int, text: str):
        row = nested.rows[row_idx]
        # Пишем в последний столбец (обычно это "поле" для текста)
        row.cells[-1].text = text

    write_block(0, desc_need)
    write_block(1, products)
    write_block(2, requirement)
    write_block(3, realization)

    # 3) Таблица задач (8-й блок) во вложенной таблице
    _fill_tasks_table(nested, res)

    # 4) Добавляем протокол ПЕРЕД ЛТ (в начало документа)
    proto_elements = _build_protocol_at_end(doc, inp, res)
    # Сейчас они в конце; вырежем и вставим в начало.
    for el in proto_elements:
        el.getparent().remove(el)
    _insert_elements_at_start(doc, proto_elements)

    # 5) Требование: весь документ — Arial 8 (включая колонтитулы)
    _force_arial_8(doc)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
